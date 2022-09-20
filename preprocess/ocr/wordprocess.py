from docx import Document
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE

doc = Document('resource\word\YOLOV3.docx')

#print(len(doc.paragraphs))
print(doc.paragraphs[0].text)
print(doc.paragraphs[1].text)
print(doc.paragraphs[2].text)
print(doc.paragraphs[3].text)

# pl = [ paragraph.text for paragraph in doc1.paragraphs]

# savepath='1.txt'


# for i in pl:
#     #print(i)
#     with open('%s' % (savepath), 'a',encoding='utf-8') as ff:
#         ff.write(i+'\n')

# #表格提取
# tables = [table for table in doc.tables]
 
# for table in tables:
#     for row in table.rows:
#         for cell in row.cells:
#             print (cell.text,end='  ')
#         print()
#     print('\n')


#页眉页脚信息

# sections = doc1.sections      #章节，一页为一章节
# sec0 = sections[0]
# head0 = sec0.header # 获取页眉对象
# head0_pars = head0.paragraphs # 获取 页眉 paragraphs
# head0_string = ''
# for par in head0_pars:
#     head0_string += par.text

# print(head0_string)

# foot0 = sec0.footer
# # <docx.section._Footer object at 0x000000000B2E3808>
# foot0_pars = foot0.paragraphs
# foot0_string = ''
# for par in foot0_pars:
#     foot0_string += par.text
# print(foot0_string)

